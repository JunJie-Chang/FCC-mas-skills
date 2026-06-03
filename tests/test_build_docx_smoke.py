"""End-to-end smoke test: JSON spec → a valid, readable .docx.

Builds from tests/fixtures/sample_spec.json, redirecting output into a
tmp dir (and disabling the ~/Downloads copy), then reopens the file with
python-docx and asserts the expected content is present.
"""
import json
from pathlib import Path

import pytest

import config
from scripts.build_docx_cli import build

FIXTURE = Path(__file__).parent / "fixtures" / "sample_spec.json"


@pytest.fixture
def isolated_output(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "OUTPUT_DIR", str(tmp_path / "output"))
    monkeypatch.setattr(config, "DISABLE_DOWNLOADS_COPY", True)
    return tmp_path


def test_build_docx_produces_valid_file(isolated_output):
    spec = json.loads(FIXTURE.read_text(encoding="utf-8"))
    out_path = build(spec)

    assert out_path.exists(), "build() did not create the .docx"
    assert out_path.suffix == ".docx"
    # stayed inside the redirected output dir
    assert str(out_path).startswith(str(isolated_output))

    # Reopen and verify it is a real, readable Word document with content.
    from docx import Document
    doc = Document(str(out_path))

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Smoke Test Report" in text   # title
    assert "Section One" in text         # heading
    assert "A paragraph of body text." in text

    # The table block rendered as an actual table with the header row.
    assert doc.tables, "expected at least one table"
    header_cells = [c.text for c in doc.tables[0].rows[0].cells]
    assert header_cells == ["A", "B"]


def test_build_docx_rejects_unknown_block(isolated_output):
    spec = {"title": "X", "task_name": "X", "intern_name": "Justin",
            "sections": [{"type": "no_such_block"}]}
    with pytest.raises(ValueError):
        build(spec)

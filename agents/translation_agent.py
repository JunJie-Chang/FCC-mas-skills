"""
agents/translation_agent.py — Translation agent using LangGraph.

Flow:
    translate → format_output

Input:  title + source + body_text (English or other language)
Output: Word doc with two sections:
    摘要 — concise Traditional Chinese summary (3–5 paragraphs)
    主文 — full paragraph-by-paragraph translation

Meta line format: YYYY-MM-DD_Source_InternName
Filename format:  YYYY.MM.DD_Title_Source_InternName.docx

Reference: Works/2026.03.11_Trump's weird Iran war_WSJ News_Justin.docx
"""
import json
import os
import sys
from datetime import date
from pathlib import Path
from typing import TypedDict, Union

from anthropic import Anthropic
from dotenv import load_dotenv
from langgraph.graph import END, StateGraph

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import config
from formatters.word_formatter import WordBuilder
from utils.file_naming import general
from utils.logger import AgentLogger

load_dotenv(Path(__file__).resolve().parents[1] / ".env")


# ── State ─────────────────────────────────────────────────────────────────────

class TranslationState(TypedDict):
    title: str
    source: str
    body_text: str
    intern_name: Union[str, list[str]]
    pub_date: str
    task_date: str
    subdir: str
    translation: dict   # {"body": [...]}
    output_path: str
    log_path: str


# ── Helpers ───────────────────────────────────────────────────────────────────

def _get_client() -> Anthropic:
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise EnvironmentError("ANTHROPIC_API_KEY not set in .env")
    return Anthropic(api_key=api_key)


def _parse_json(text: str):
    import re
    original = text
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()
    match = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", text)
    if match:
        text = match.group(1)
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        print(f"\n[DEBUG] JSON parse failed: {e}")
        print(f"[DEBUG] Raw LLM response:\n{original[:2000]}")
        raise


def _read_body_file(path: str) -> str:
    """Read body text from .txt, .docx, or an image/PDF (OCR via Claude Vision)."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"找不到檔案：{path}")
    suffix = p.suffix.lower()
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(p))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    if suffix in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".pdf"}:
        from utils.ocr import extract_text
        print(f"OCR 中：{p.name}")
        text = extract_text(str(p))
        print(f"  OCR 完成，擷取 {len(text)} 字元")
        return text
    return p.read_text(encoding="utf-8")


# ── Node 1: translate ─────────────────────────────────────────────────────────

def translate(state: TranslationState) -> dict:
    """Translate body text into Traditional Chinese; produce summary + full body."""
    client = _get_client()

    prompt = f"""你是一位專業的繁體中文翻譯，服務於香港金融顧問公司。
請將以下文章逐段翻譯成繁體中文：
   - 保留原文段落結構（每段對應一個 JSON 元素）
   - 專有名詞保留英文，人名譯音加括號標示（如「川普（Trump）」）
   - 語氣自然流暢，符合繁體中文書面語習慣

標題：{state['title']}
來源：{state['source']}

原文：
{state['body_text']}

回傳純 JSON，不要 markdown 包裝，格式如下：
{{
  "body": [
    "第一段翻譯",
    "第二段翻譯"
  ]
}}"""

    message = client.messages.create(
        model=config.LLM_FAST,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )

    from utils.cost_tracker import tracker
    tracker.record_claude(config.LLM_FAST, message.usage.input_tokens, message.usage.output_tokens)

    raw = next((b.text for b in message.content if hasattr(b, "text")), "")
    return {"translation": _parse_json(raw)}


# ── Node 2: format_output ─────────────────────────────────────────────────────

def format_output(state: TranslationState) -> dict:
    translation = state["translation"]
    intern = state["intern_name"]
    pub_date = state.get("pub_date") or date.today().strftime("%Y-%m-%d")
    task_date = state.get("task_date") or date.today().strftime("%Y-%m-%d")
    subdir = state.get("subdir", "adhoc")
    source = state["source"]
    title = state["title"]

    if isinstance(intern, list):
        intern_str = ", ".join(intern)
    else:
        intern_str = intern or config.DEFAULT_INTERN_NAME

    # Meta: two lines — pub date + source, then work date + intern
    meta_text = f"{pub_date}_{source}\n{task_date}_{intern_str}"

    builder = WordBuilder(title, task_date, intern, meta_text=meta_text)
    builder.add_blank_line()
    for para in translation.get("body", []):
        builder.add_paragraph(para)

    dot_date = task_date.replace("-", ".")
    # Filename: YYYY.MM.DD_Title_Source_InternName.docx
    filename = general(f"{title}_{source}", intern, dot_date)
    output_path = builder.save(filename, subdir=subdir)

    logger = AgentLogger(
        "translation_agent",
        f"{title} [{source}]",
        intern,
    )
    log_path = logger.save(output_path)

    return {"output_path": str(output_path), "log_path": str(log_path)}


# ── Graph ─────────────────────────────────────────────────────────────────────

def build_graph():
    graph = StateGraph(TranslationState)
    graph.add_node("translate", translate)
    graph.add_node("format_output", format_output)

    graph.set_entry_point("translate")
    graph.add_edge("translate", "format_output")
    graph.add_edge("format_output", END)

    return graph.compile()


# ── Public entry point ────────────────────────────────────────────────────────

def run(
    title: str,
    source: str,
    body_text: str,
    intern_name: Union[str, list[str]] = None,
    pub_date: str = None,
    task_date: str = None,
    subdir: str = "daily",
    # router compat: ignored when title/body_text are provided directly
    task_instruction: str = None,
) -> dict:
    """
    Run the translation agent.

    Args:
        title:       Original article title (English/source language).
        source:      Publication name, e.g. "WSJ News", "FT".
        body_text:   Full article body to translate.
        intern_name: Intern name(s) for file naming.
        task_date:   Date string YYYY-MM-DD; defaults to today.
        subdir:      Output subfolder — 'daily', 'weekly', or 'adhoc'.

    Returns:
        dict with 'output_path' and 'log_path'.
    """
    app = build_graph()

    final_state = app.invoke({
        "title":       title,
        "source":      source,
        "body_text":   body_text,
        "intern_name": intern_name or config.DEFAULT_INTERN_NAME,
        "pub_date":    pub_date or date.today().strftime("%Y-%m-%d"),
        "task_date":   task_date or date.today().strftime("%Y-%m-%d"),
        "subdir":      subdir,
        "translation": {},
        "output_path": "",
        "log_path":    "",
    })

    return {
        "output_path": final_state["output_path"],
        "log_path":    final_state["log_path"],
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Translation agent — translate to Traditional Chinese")
    parser.add_argument("--title",     required=True, help="Original article title")
    parser.add_argument("--source",    required=True, help="Publication source, e.g. 'WSJ News'")
    parser.add_argument("--body",      default=None,  help="Article body text (inline)")
    parser.add_argument("--body-file", default=None,  dest="body_file",
                        help="Path to .txt or .docx file containing the article body")
    parser.add_argument("--intern",    default=None)
    parser.add_argument("--pub-date",  default=None,  dest="pub_date",
                        help="Publication date YYYY-MM-DD (shown on meta line 1)")
    parser.add_argument("--date",      default=None,  help="Work date YYYY-MM-DD (shown on meta line 2)")
    parser.add_argument("--subdir",    default="daily", choices=["daily", "weekly", "adhoc"])
    args = parser.parse_args()

    # Resolve body text
    if args.body_file:
        body_text = _read_body_file(args.body_file)
    elif args.body:
        body_text = args.body
    else:
        parser.error("Provide --body <text> or --body-file <path>")

    intern = (
        [n.strip() for n in args.intern.split(",")]
        if args.intern and "," in args.intern
        else args.intern
    )

    print(f"翻譯中：{args.title} [{args.source}]")
    result = run(
        title=args.title,
        source=args.source,
        body_text=body_text,
        intern_name=intern,
        pub_date=args.pub_date,
        task_date=args.date,
        subdir=args.subdir,
    )
    print(f"  Output: {result['output_path']}")
    print(f"  Log:    {result['log_path']}")

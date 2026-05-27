"""
formatters/word_formatter.py — Unified Word document builder.

All agents use WordBuilder to produce .docx output.
After saving, the file is automatically copied to ~/Downloads.

Observed format from Works/ directory:
  - Font:    微軟正黑體 on every run (ascii + eastAsia + hAnsi + cs)
  - Title:   14pt, Bold (Para[0])
  - Meta:    14pt, Normal — "YYYY-MM-DD, InternName" (Para[1])
  - Heading: 14pt, Bold
  - Body:    14pt, Normal
  - Page:    A4, margins top/bottom=2.5cm, left/right=3.2cm
  - Spacing: 1.15 multiple line spacing
  - Border:  single-line page border on all four sides (sz=12, space=24)
  - Header:  "Private & Confidential" — right-aligned, 11pt Bold
"""
import re
import shutil
from datetime import date
from pathlib import Path
from typing import Union

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import config

_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_CITATION_RE = re.compile(r"\[(\d+)\]")

_COMMENTS_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── Internal helpers ──────────────────────────────────────────────────────────

def _apply_font(run, size_pt: float = None, bold: bool = None) -> None:
    """
    Set 微軟正黑體 on a run for all four font slots (ascii, hAnsi, eastAsia, cs)
    so that both Latin and CJK characters render correctly.
    """
    font_name = config.FONT_NAME
    run.font.name = font_name  # sets ascii + hAnsi

    # Explicitly set eastAsia and cs (python-docx's font.name setter misses these)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _apply_para_spacing(para) -> None:
    """Apply 1.15 multiple line spacing (no extra space-before/after by default)."""
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = config.LINE_SPACING


def _add_hyperlink_run(para, text: str, url: str, size_pt: float) -> None:
    """Insert a styled hyperlink run (blue, underlined) into an existing paragraph."""
    r_id = para.part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(f"{{{_R_NS}}}id", r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), config.FONT_NAME)
    rPr.append(rFonts)

    half_pts = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), half_pts)
        rPr.append(el)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def _add_runs_with_citations(
    para, text: str, references: list[dict], size_pt: float, bold: bool = False
) -> None:
    """
    Add runs to para, turning [N] markers into clickable hyperlinks.
    Falls back to a plain run when references is empty or a citation has no URL.
    """
    if not references:
        run = para.add_run(text)
        _apply_font(run, size_pt=size_pt, bold=bold)
        return

    ref_map = {r["num"]: r for r in references}
    parts = _CITATION_RE.split(text)  # [text, num, text, num, ...]

    i = 0
    while i < len(parts):
        chunk = parts[i]
        if chunk:
            run = para.add_run(chunk)
            _apply_font(run, size_pt=size_pt, bold=bold)
        i += 1
        if i < len(parts):
            num = int(parts[i])
            ref = ref_map.get(num)
            if ref and ref.get("url"):
                _add_hyperlink_run(para, f"[{num}]", ref["url"], size_pt)
            else:
                run = para.add_run(f"[{num}]")
                _apply_font(run, size_pt=size_pt, bold=bold)
            i += 1


# ── WordBuilder ───────────────────────────────────────────────────────────────

class WordBuilder:
    """
    Fluent builder for Word documents in FCC format.

    Usage:
        doc = (
            WordBuilder("Tesla 自動駕駛調查", "2026.04.07", "Justin")
            .add_heading("公司背景")
            .add_paragraph("Tesla 成立於 2003 年...")
            .add_bullet("Autopilot 累積里程超過 5 億英里")
        )
        path = doc.save("2026.04.07_Tesla 自動駕駛調查_Justin.docx", subdir="daily")
    """

    def __init__(
        self,
        title: str,
        task_date: Union[str, date] = None,
        intern_name: Union[str, list[str]] = None,
        meta_text: str = None,
    ):
        self._doc = Document()
        self._comment_counter = 0
        self._comments_part_ref = None
        self._setup_page()
        self._add_header(title, task_date, intern_name, meta_text=meta_text)

    # ── Setup ─────────────────────────────────────────────────────────────────

    def _setup_page(self) -> None:
        section = self._doc.sections[0]
        section.page_width = Cm(config.PAGE_WIDTH_CM)
        section.page_height = Cm(config.PAGE_HEIGHT_CM)
        section.top_margin = Cm(config.MARGIN_TOP_CM)
        section.bottom_margin = Cm(config.MARGIN_BOTTOM_CM)
        section.left_margin = Cm(config.MARGIN_LEFT_CM)
        section.right_margin = Cm(config.MARGIN_RIGHT_CM)
        self._add_page_border(section)
        self._add_confidential_header(section)
        self._add_page_number_footer(section)

    @staticmethod
    def _add_page_border(section) -> None:
        """Add single-line page border on all four sides (matches Works/ files)."""
        sectPr = section._sectPr
        # Remove existing pgBorders if present
        for existing in sectPr.findall(qn("w:pgBorders")):
            sectPr.remove(existing)
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "12")
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), "auto")
            pgBorders.append(border)
        sectPr.append(pgBorders)

    @staticmethod
    def _add_page_number_footer(section) -> None:
        """Add centered page number in the footer using a PAGE field."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        footer = section.footer
        for para in footer.paragraphs:
            para._element.getparent().remove(para._element)
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        _apply_font(run, size_pt=11, bold=False)
        # Insert PAGE field: <w:fldChar w:fldCharType="begin"/> ... PAGE ... end
        fldBegin = OxmlElement("w:fldChar")
        fldBegin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldSep = OxmlElement("w:fldChar")
        fldSep.set(qn("w:fldCharType"), "separate")
        fldEnd = OxmlElement("w:fldChar")
        fldEnd.set(qn("w:fldCharType"), "end")
        run._r.append(fldBegin)
        run._r.append(instrText)
        run._r.append(fldSep)
        run._r.append(fldEnd)

    @staticmethod
    def _add_confidential_header(section) -> None:
        """
        Add 'Private & Confidential' to the page header, positioned flush
        against the right page border (matches Works/ reference files).

        Uses a negative right indent (-1050 twips) identical to the original,
        which pushes the text beyond the text area margin to hug the border.
        """
        header = section.header
        for para in header.paragraphs:
            para._element.getparent().remove(para._element)
        p = header.add_paragraph()

        # Set paragraph properties: right-aligned + negative right indent
        pPr = p._p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:right"), "-1050")
        ind.set(qn("w:firstLine"), "440")
        pPr.append(ind)

        run = p.add_run("Private & Confidential")
        _apply_font(run, size_pt=11, bold=True)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    def _add_header(
        self,
        title: str,
        task_date: Union[str, date] = None,
        intern_name: Union[str, list[str]] = None,
        meta_text: str = None,
    ) -> None:
        # Para 0 — Title
        p_title = self._doc.add_paragraph()
        run = p_title.add_run(title)
        _apply_font(run, size_pt=config.FONT_SIZE_TITLE_PT, bold=True)
        _apply_para_spacing(p_title)

        # Para 1 — Meta line (custom or default)
        if meta_text:
            for line in meta_text.split("\n"):
                p_meta = self._doc.add_paragraph()
                run_meta = p_meta.add_run(line)
                _apply_font(run_meta, size_pt=config.FONT_SIZE_PT)
                _apply_para_spacing(p_meta)
        elif task_date or intern_name:
            if isinstance(task_date, date):
                date_str = task_date.strftime("%Y-%m-%d")
            else:
                date_str = str(task_date) if task_date else date.today().strftime("%Y-%m-%d")

            if isinstance(intern_name, list):
                intern_str = ", ".join(intern_name)
            else:
                intern_str = intern_name or config.DEFAULT_INTERN_NAME

            p_meta = self._doc.add_paragraph()
            run_meta = p_meta.add_run(f"{date_str}, {intern_str}")
            _apply_font(run_meta, size_pt=config.FONT_SIZE_PT)
            _apply_para_spacing(p_meta)

    # ── Content blocks ────────────────────────────────────────────────────────

    def add_heading(self, text: str) -> "WordBuilder":
        """Bold section heading (same size as body, distinguished by bold)."""
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=True)
        _apply_para_spacing(p)
        return self

    def add_paragraph(self, text: str, references: list[dict] = None) -> "WordBuilder":
        """Normal body paragraph. Pass references to hyperlink [N] citation markers."""
        p = self._doc.add_paragraph()
        _add_runs_with_citations(p, text, references or [], config.FONT_SIZE_PT, bold=False)
        _apply_para_spacing(p)
        return self

    def add_bullet(self, text: str, references: list[dict] = None) -> "WordBuilder":
        """Single bullet list item. Pass references to hyperlink [N] citation markers."""
        p = self._doc.add_paragraph(style="List Bullet")
        _add_runs_with_citations(p, text, references or [], config.FONT_SIZE_PT, bold=False)
        return self

    def add_references(self, references: list[dict]) -> "WordBuilder":
        """Render a '參考來源' section with numbered clickable hyperlinks."""
        if not references:
            return self
        self.add_blank_line()
        self.add_heading("參考來源")
        for ref in references:
            p = self._doc.add_paragraph(style="List Bullet")
            label_run = p.add_run(f"[{ref['num']}]  {ref.get('title', '')}  ")
            _apply_font(label_run, size_pt=config.FONT_SIZE_PT, bold=False)
            url = ref.get("url", "")
            if url:
                _add_hyperlink_run(p, url, url, config.FONT_SIZE_PT)
        return self

    def add_red_heading(self, text: str) -> "WordBuilder":
        """Bold red heading — used for podcast article titles."""
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=True)
        run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
        _apply_para_spacing(p)
        return self

    def add_red_underline_title_with_subtitle(
        self, title: str, subtitle: str = "",
    ) -> "WordBuilder":
        """
        Podcast article title block — one paragraph with two runs separated by
        a soft line break (shift+enter):
          line 1: title  — Bold + red + underlined
          line 2: subtitle — Normal black (e.g. "2025.10.17_Justin_中央社_林宏翰")
        """
        p = self._doc.add_paragraph()
        title_run = p.add_run(title)
        _apply_font(title_run, size_pt=config.FONT_SIZE_PT, bold=True)
        title_run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
        title_run.font.underline = True

        if subtitle:
            # Soft line break (w:br) — equivalent to shift+enter, stays in same paragraph
            br = OxmlElement("w:br")
            title_run._r.append(br)
            sub_run = p.add_run(subtitle)
            _apply_font(sub_run, size_pt=config.FONT_SIZE_PT, bold=False)

        _apply_para_spacing(p)
        return self

    # ── Word comments ─────────────────────────────────────────────────────────

    def _get_or_create_comments_part(self):
        """Return the comments.xml part, creating it if absent."""
        if self._comments_part_ref is not None:
            return self._comments_part_ref

        from lxml import etree
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        doc_part = self._doc.part
        for rel in doc_part.rels.values():
            if rel.reltype == _COMMENTS_REL:
                self._comments_part_ref = rel.target_part
                return self._comments_part_ref

        root = etree.Element(f"{{{_W_NS}}}comments", nsmap={"w": _W_NS})
        xml_bytes = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        new_part = Part(
            PackURI("/word/comments.xml"),
            _COMMENTS_CT,
            xml_bytes,
            doc_part.package,
        )
        doc_part.relate_to(new_part, _COMMENTS_REL)
        self._comments_part_ref = new_part
        return new_part

    def _attach_comment(self, para, comment_text: str) -> None:
        """Attach a Word comment bubble to an existing paragraph."""
        from lxml import etree
        from datetime import datetime, timezone

        cid = self._comment_counter
        self._comment_counter += 1
        cid_str = str(cid)

        def w(tag):
            return f"{{{_W_NS}}}{tag}"

        # ── 1. Add entry to comments.xml ──────────────────────────────────────
        cp = self._get_or_create_comments_part()
        comments_root = etree.fromstring(cp.blob)

        comment_el = etree.SubElement(comments_root, w("comment"))
        comment_el.set(w("id"), cid_str)
        comment_el.set(w("author"), "FCC MAS")
        comment_el.set(
            w("date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        )
        comment_el.set(w("initials"), "F")

        p_el = etree.SubElement(comment_el, w("p"))
        r_el = etree.SubElement(p_el, w("r"))
        t_el = etree.SubElement(r_el, w("t"))
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = comment_text

        cp._blob = etree.tostring(
            comments_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        # ── 2. Insert markers into the paragraph ──────────────────────────────
        p_xml = para._p

        rng_start = OxmlElement("w:commentRangeStart")
        rng_start.set(qn("w:id"), cid_str)
        first_run = p_xml.find(qn("w:r"))
        if first_run is not None:
            first_run.addprevious(rng_start)
        else:
            p_xml.insert(0, rng_start)

        rng_end = OxmlElement("w:commentRangeEnd")
        rng_end.set(qn("w:id"), cid_str)
        p_xml.append(rng_end)

        ref_r = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        ref_sty = OxmlElement("w:rStyle")
        ref_sty.set(qn("w:val"), "CommentReference")
        ref_rpr.append(ref_sty)
        ref_r.append(ref_rpr)
        ref_ref = OxmlElement("w:commentReference")
        ref_ref.set(qn("w:id"), cid_str)
        ref_r.append(ref_ref)
        p_xml.append(ref_r)

    def add_red_heading_with_comment(
        self, text: str, comment_text: str
    ) -> "WordBuilder":
        """Bold red heading with a Word comment bubble containing source metadata."""
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=True)
        run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
        _apply_para_spacing(p)
        if comment_text:
            self._attach_comment(p, comment_text)
        return self

    def add_bracket_heading(self, text: str) -> "WordBuilder":
        """Section heading in 【】brackets, fully bold — e.g. 【Meeting Executive Summary】."""
        p = self._doc.add_paragraph()
        run = p.add_run(f"【{text}】")
        _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=True)
        _apply_para_spacing(p)
        return self

    def add_topic_heading(self, text: str) -> "WordBuilder":
        """Discussion sub-group heading — normal weight (not bold)."""
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=False)
        _apply_para_spacing(p)
        return self

    def add_keyed_paragraph(self, key: str, text: str) -> "WordBuilder":
        """Bold key + normal explanation in one paragraph: 'Key: explanation'.
        Used for discussion points and action items."""
        p = self._doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        _apply_font(run_key, size_pt=config.FONT_SIZE_PT, bold=True)
        run_text = p.add_run(text)
        _apply_font(run_text, size_pt=config.FONT_SIZE_PT, bold=False)
        _apply_para_spacing(p)
        return self

    def add_keyed_info(self, key: str, text: str) -> "WordBuilder":
        """Bold key + normal text for info block: 'Key: value'. No extra spacing."""
        p = self._doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        _apply_font(run_key, size_pt=config.FONT_SIZE_PT, bold=True)
        run_text = p.add_run(text)
        _apply_font(run_text, size_pt=config.FONT_SIZE_PT, bold=False)
        _apply_para_spacing(p)
        return self

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
    ) -> "WordBuilder":
        """Insert a table with a bold header row and single-line black borders on all cells."""
        table = self._doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"   # built-in style: single-line borders on all sides + insideH/V
        # Header row
        for i, h in enumerate(headers):
            run = table.rows[0].cells[i].paragraphs[0].add_run(h)
            _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=True)
        # Data rows
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                run = table.rows[ri + 1].cells[ci].paragraphs[0].add_run(str(val))
                _apply_font(run, size_pt=config.FONT_SIZE_PT, bold=False)
        return self

    def add_blank_line(self) -> "WordBuilder":
        """Insert an empty paragraph as a visual spacer."""
        p = self._doc.add_paragraph()
        _apply_para_spacing(p)
        return self

    # ── Save ──────────────────────────────────────────────────────────────────

    def save(
        self,
        filename: str,
        subdir: str = "adhoc",
        output_dir: str = None,
    ) -> Path:
        """
        Save the document and copy it to ~/Downloads.

        Args:
            filename:   Final filename including extension, e.g.
                        '2026.04.07_Tesla 自動駕駛調查_Justin.docx'
                        Use utils.file_naming to generate this.
            subdir:     Subfolder under output/ — 'daily', 'weekly', or 'adhoc'.
            output_dir: Override the default output directory from config.

        Returns:
            Path to the saved file inside output/.
        """
        base = Path(output_dir or config.OUTPUT_DIR) / subdir
        base.mkdir(parents=True, exist_ok=True)
        dest = base / filename
        self._doc.save(str(dest))

        # Auto-copy to ~/Downloads (skipped in server mode — interns
        # fetch via /api/jobs/{id}/download, not from the host's disk).
        if not config.DISABLE_DOWNLOADS_COPY:
            dl_path = Path(config.DOWNLOADS_DIR) / filename
            shutil.copy2(dest, dl_path)

        return dest
